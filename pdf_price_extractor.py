# pdf_price_extractor.py
# PDF/ZIP Price Extraction Module for Moroccan Public Tenders
# This module integrates with automate marchepublics - Copy.py

import zipfile
import os
import re
import json
import subprocess
import tempfile
import time
from io import BytesIO
from dotenv import load_dotenv
from google import genai
from google.genai import types

# 🔑 Load .env independently — this module can be exec'd by the launcher
# BEFORE the main script runs its own load_dotenv(), so we cannot rely on
# that call having already happened. load_dotenv() is safe to call multiple
# times (it's a no-op if variables are already set / .env already read).
load_dotenv()

def safe_math_eval(formula_str):
    """Grand-Master Safe Evaluator (Standard 64-bit Precision)"""
    import re
    if not formula_str or not isinstance(formula_str, str): return None
    
    # 🇲🇦 MOROCCAN COMMA SHIELD: Convert commas to dots before validation
    # This prevents failure if Gemini reads "17,92" from a French PDF
    formula_str = formula_str.replace(',', '.')
    
    # 🛡️ SECURITY SHIELD: Only allow digits and basic math operators
    if not re.match(r'^[0-9.+\-*/() ]+$', formula_str): 
        print(f"⚠️ Formula rejected (illegal characters): {formula_str}")
        return None
    try:
        # Standard calculation with floating point precision
        result = eval(formula_str, {"__builtins__": {}}, {})
        return float(result)
    except Exception as e:
        print(f"⚠️ Math Execution Error: {e}")
        return None


# 🧮 SMIG UNIT-PRICE CALCULATOR
# Centralised helper so the "17.92 DH/heure SMIG" conversion logic lives in
# ONE place instead of being duplicated. Used as the role-matched-hours
# fallback when Gemini reports hours/agents but no explicit price.
def compute_smig_unit_price(unit_type, effective_hours=8.0, smig_hourly=17.92, days_per_month=26):
    """
    كتحسب الثمن الوحدوي (prix unitaire) اعتمادا على SMIG بالساعة ونوع الوحدة.

    - "heure" -> السميغ بالساعة مباشرة (مثلا 17.92 DH)
    - "jour"  -> السميغ × عدد الساعات فالنهار (مثلا 17.92 × 4 = 71.68 DH)
    - "mois"  -> السميغ × الساعات × 26 يوم (مثلا 17.92 × 4 × 26 = 1863.68 DH)

    Args:
        unit_type (str): نص فيه "heure" / "jour" / "mois" (case-insensitive,
                          partial match زي u_type ديال الكود الأصلي).
        effective_hours (float): عدد ساعات الخدمة فالنهار (h). Default 8.0.
        smig_hourly (float): قيمة SMIG بالساعة. Default 17.92 DH.
        days_per_month (int): عدد أيام العمل فالشهر. Default 26.

    Returns:
        float | None: الثمن المحسوب مقرب لـ 2 أرقام عشرية، ولا None إلا
                       ماكانش نوع الوحدة معروف (heure/jour/mois).
    """
    u_type = str(unit_type).lower() if unit_type else ""
    h = float(effective_hours) if effective_hours else 8.0

    if "jour" in u_type:
        computed_price = round(smig_hourly * h, 2)
        print(f"🧮 SMIG Calc: {smig_hourly} × {h}h = {computed_price} DH/jour")
        return computed_price

    elif "mois" in u_type:
        computed_price = round(smig_hourly * h * days_per_month, 2)
        print(f"🧮 SMIG Calc: {smig_hourly} × {h}h × {days_per_month}j = {computed_price} DH/mois")
        return computed_price

    elif "heure" in u_type:
        computed_price = round(smig_hourly, 2)
        print(f"🧮 SMIG Calc: {computed_price} DH/heure (direct)")
        return computed_price

    else:
        print(f"⚠️ SMIG Calc: unit type '{unit_type}' not recognized (heure/jour/mois).")
        return None


# --- CONFIGURATION ---
DOWNLOAD_PATH = r"C:\Users\hp\Downloads\MarchePublics"
LIBREOFFICE_PATH = r"C:\Program Files\LibreOffice\program\soffice.exe"
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if not GEMINI_API_KEY:
    raise RuntimeError(
        "GEMINI_API_KEY not found in environment. Make sure it's set in your .env "
        "file (the main script loads it via load_dotenv() before importing this module)."
    )
gemini_client = genai.Client(api_key=GEMINI_API_KEY)

# Ensure download directory exists
os.makedirs(DOWNLOAD_PATH, exist_ok=True)

# --- VISION PRICE EXTRACTION (GEMINI 3 PRO) ---
def extract_price_with_gemini_vision(file_bytes, unit_type, hours_per_day=None, role_description=None):
    """
    Sends the FULL PDF to Gemini 3 Pro (Vision Mode).
    Uses 'Hierarchy of Truth' prompt to extract price from any layout (scanned/digital).
    """
    u_type = str(unit_type).lower() if unit_type else ""

    total_hours_block = f"""
[TOTAL WORK-HOURS FOR A LUMP-SUM "FORFAIT" LINE]
Unit Type Context is "{unit_type}". If this is "forfait" (a single lump-sum line
covering the WHOLE contract, e.g. "Lot unique", qty=1, no per-role breakdown),
Moroccan gardiennage/surveillance CPS documents very often state the TOTAL
work-hours for the ENTIRE contract directly in prose, near the object/duration
clause — e.g. "Délai d'exécution... 2112 Heures pendant 88 jours", or "soit un
total de X heures". This is usually the SUM of hours worked by ALL agents
combined (all shifts, all agents) over the full contract duration — NOT a
per-agent or per-day figure.
- Search the document (Article 1 "Objet"/"Délai d'exécution" and Article 2
  "Effectifs" are the most common locations) for this total-hours statement.
- If found, extract the raw number into "total_hours_all_agents" (e.g. 2112).
  ⚠️ Report this field WHENEVER you find such a statement — independent of
  whether a price/price-breakdown table exists. This is DATA you observed
  (a number written in the text), not a price calculation, so it is always
  safe to report even if "price" ends up null.
- A staffing table (e.g. "Effectif | jours ouvrables | jours fériés" with rows
  like Chef d'équipe, agents, télésurveillance...) is a supporting detail —
  if it has its own "heures" total row, that number should match the prose
  statement; use it to confirm, not as a substitute if the two only look
  similar. Prefer the explicit prose total when both are present.
- If no such total-hours statement exists anywhere, set
  "total_hours_all_agents": null.
"""

    role_context_block = ""
    if role_description:
        role_context_block = f"""
[TARGET ROLE / LINE ITEM]
The price you must find applies SPECIFICALLY to this role/line item (as written on the tender's website table):
"{role_description}"

⚠️ CRITICAL: Many Moroccan gardiennage/surveillance tenders list a staffing table
(often under a heading like "Article 3. Effectif, horaires et organisation" or
"Poste | Nombre d'agents | Horaire") with SEVERAL DIFFERENT ROLES, each having its
OWN number of agents and its OWN working hours per shift (e.g. Chef d'équipe 12h,
Agent jour 12h, Agent soir 12h — these are often NOT the default 8h).
- Find the ROW in that table that matches the TARGET ROLE above (match by role name:
  "Chef d'équipe", "Agent ... jour", "Agent ... soir", "Maître-chien", etc.)
- Extract the EXACT hours-per-shift for THAT row into "matched_hours_per_day".
  ⚠️ Fill this field WHENEVER you find a matching staffing/schedule table row —
  this is independent from whether a PRICE or price-breakdown table exists
  elsewhere. Even if you cannot find any price at all in the document (so
  "price" will be null), still report "matched_hours_per_day" if you found it.
  This is just DATA you observed, not a price calculation, so it's always safe
  to report.
- Do NOT average across all roles and do NOT reuse another role's hours.
- If the table has a per-role UNIT PRICE column already filled in for this exact
  role, prefer that price (Priority 3) over any general/global calculation.
- If you cannot find ANY table row matching this specific role anywhere in the
  document, set "matched_hours_per_day": null and proceed with the normal
  hierarchy below.

[TOTAL AGENT COUNT — ALWAYS ATTEMPT THIS, SEPARATELY FROM THE ABOVE]
Independently of the role-matching above, search the ENTIRE document for the
TOTAL number of agents/personnel required for the WHOLE contract (not just one
role/shift). Look for phrases like "par X agents qualifiés", "effectif de X
agents", "un effectif de X personnes", "X jardiniers qualifiés présents à
chaque intervention", or a "Total" row at the bottom of a staffing/schedule
table (e.g. "Total: 13 / 8 / 8" per shift — in that case sum the distinct
people, or take the largest simultaneous total, whichever the document makes
clearest).
- Report this as "matched_total_agents" (an integer).
- This is independent from "matched_hours_per_day" — fill both if you can find
  them, fill only the one you find if you can't find the other, and set
  "matched_total_agents": null if no total agent count is stated anywhere.
- This is just DATA you observed, not a price — always safe to report even
  when "price" ends up null.
"""

    prompt = f"""[ROLE]
You are a Compliance Officer & Data Extraction Engine for Moroccan public tenders.
Your Goal: Understand the document's meaning perfectly—regardless of layout or phrasing—and identify the MINIMUM COMPLIANT BID PRICE (Prix Unitaire HT).

[INPUT_CONTEXT]
- File detected as: FULL PDF DOCUMENT (Vision Mode)
- Unit Type Context: {unit_type}
{role_context_block}

[HIERARCHY_OF_TRUTH] (Follow Order 1 -> 3 EXACTLY)
You must search for the price in this specific order of authority. Once a valid price is found, STOP and return it.

## PRIORITY 1: EXPLICIT TEXT WARNINGS (Overrules Tables)
Search the text contexts (outside the tables) for mandatory pricing thresholds.
- Look for: "toute offre inférieure à X sera écartée"
- Look for: "le prix ne doit pas être inférieur à X"
- Look for: "est arrêté à la somme de X"
- Look for: "Minimum: X DH"
- Look for (ARABIC): "يقل عن" (less than), "الحد الأدنى" (minimum), "سعر ادنى" (min price), "مبلغ أدنى" (min amount)
-> IF FOUND: Return X. Set "found_in_text_warning": true.

⚠️ ROLE-SPECIFICITY GUARD (applies when [TARGET ROLE / LINE ITEM] is provided above):
- ONLY use this priority if the warning/minimum applies DIRECTLY to the target role
  itself (e.g. "le prix du Chef d'équipe ne doit pas être inférieur à X").
- Do NOT use a GLOBAL/PROJECT-WIDE figure — such as the overall "estimation du
  maître d'ouvrage" (total contract value) combined with the eviction rule
  ("offre inférieure de plus de 20% est écartée") — as if it were a per-role price.
  That total covers ALL roles combined over the full contract; dividing it evenly
  across agents/months produces a WRONG flat rate that ignores that different
  roles (chef d'équipe, night shifts, maître-chien, etc.) are paid differently.
- If the only "warning" you can find is this kind of global estimation math, treat
  it as NOT FOUND for Priority 1 — set "found_in_text_warning": false and continue
  to Priority 2/3, using the role-matched staffing table row instead (see above).

## PRIORITY 2: TEMPLATE CALCULATION (Empty Tables)
 If columns 12, 13, 14 exist but are empty or contain formulas, YOU MUST CALCULATE THE PRICE.
 ⚠️ CRITICAL CONDITION: ONLY CALCULATE IF A SPECIFIC PRICING BREAKDOWN TABLE EXISTS (e.g. columns for 'Salaire', 'CNSS', 'Charges', 'Total').
 - IF NO SUCH TABLE EXISTS (e.g. only text description or simple list of articles), **DO NOT CALCULATE**. Return `null`.

 Analyze the table's implicit formula ("Follow what the table has"):

 Method A (CNSS on Base - Common):
 - Base SMIG: 17.92
 - Congés (5.77%) + Fériés (3.85%) = 9.62% of 17.92 = 1.72
 - CNSS (21.09%) is calculated on BASE SMIG (17.92) = 3.78
 - Total = 17.92 + 1.72 + 3.78 ~= 23.42

Method B (CNSS on Gross - safer/higher):
- Base SMIG: 17.92
- Gross Salary = 17.92 + 1.72 = 19.64
- CNSS (21.09%) is calculated on GROSS (19.64) = 4.14
- Total = 19.64 + 4.14 ~= 23.78

INSTRUCTION: Check the column headers or any example rows.
- If the table implies CNSS is on Base (or if ambiguous/standard), USE METHOD A (Target ~23.42).
- If the table implies CNSS is on Gross, USE METHOD B (Target ~23.78).
- DO NOT add Assurance/Marge. Return the RAW CALCULATED SUM. Python will add the margins.
- IMPORTANT: If calculation results in 3 decimals (e.g. 28.625), ALWAYS ROUND DOWN (TRUNCATE) to 2 decimals (e.g. 28.62) to be competitive.

-> Return the CALCULATED VALUE as "price". Set "is_calculated": true. AND Set "found_price_is_hourly": true.

## PRIORITY 3: FINAL PRICE COLUMN (Filled Tables)
If no text warning exists, find the final price in the table columns.
Search Order:
1. "Prix unitaire journalier en dirhams hors TVA" / "Prix unitaire HT"
2. "Total heur HT" / "Total jour HT" (Column 15 style)
3. "TOTAL (IV)" / "TOTAL IV"
4. "TOTAL (3)" / "TOTAL III"
5. "TOTAL (1)" (Only if followed by separate 'Marge' column)

## CRITICAL UNIT IDENTIFICATION:
1. IF the price found is explicitly per HOUR (e.g. "20 DH/h"):
   - Set "found_price_is_hourly": true.
   - Return the RAW value (e.g. 20). DO NOT CALCULATE.
2. IF the price found is explicitly per DAY (e.g. "160 DH/j"):
   - Set "found_price_is_hourly": false.
   - Return the value as is.

[VALIDATION_RULES]
1. Price must be a raw numeric value without currency symbols (e.g. 190.82, never "190,82 DH").
2. Price must be strictly rounded to 2 decimal places (e.g. 190.82, never 190.82397).
3. Price must be > 0.
4. CHECK FOR COST COMPONENTS (CONTEXT-AWARE):
   - Scan table for 'Assurance', 'Assur', 'AT', or 'RC'. IF AND ONLY IF such a column EXISTS but the cell is empty or zero -> set 'missing_assurance': true. If the column is NOT PRESENT -> set 'missing_assurance': false.
   - Scan table for 'Charge', 'Prestations', 'Cotisations', 'Charges de fonctionnement', 'Tenues', 'Matériel', or 'Autre frais'. IF AND ONLY IF such a column EXISTS but the cell is empty or zero -> set 'missing_charge': true. If the column is NOT PRESENT -> set 'missing_charge': false.
   - Scan table for 'Marge', 'Marge bénéficiaire', or 'Bénéfice'. IF AND ONLY IF such a column EXISTS but the cell is empty or zero -> set 'missing_margin': true. If the column is NOT PRESENT -> set 'missing_margin': false.

[EXAMPLES]

## Ex 1: Text Warning (Overrides Table)
Input: "Table: [Empty] | Text: 'Toute offre inférieure à 190.82 DH sera rejetée'"
Output: {{"price": 190.82, "found_in_text_warning": true, "confidence": "high"}}

## Ex 2: Standard Table
Input: "Prix unitaire journalier: 80.53"
Output: {{"price": 80.53, "found_in_text_warning": false, "confidence": "high"}}

## Ex 3: Template
Input: "Smig (17.92) | Congés (5.77%) | Fériés (3.85%)"
Output: {{"price": 19.64, "calculation_formula": "17.92 + (17.92 * 0.0577) + (17.92 * 0.0385)", "is_calculated": true, "confidence": "high"}}

## Ex 4: Missing Columns
Input: "Total(1)=185.33 | Assur=[Empty] | charge=[Empty] | Marge=[Empty]"
Output: {{"price": 185.33, "missing_assurance": true, "missing_charge": true, "missing_margin": true, "confidence": "high"}}

## Ex 5: Hourly Price (Python will convert)
Input: "Prix unitaire horaire: 20.00"
Output: {{"price": 20.00, "found_price_is_hourly": true, "confidence": "high"}}

[OUTPUT_SCHEMA]
Return ONLY valid JSON:
{{
    "price": "number|null",
    "calculation_formula": "string|null",
    "found_in_text_warning": "boolean",
    "is_calculated": "boolean",
    "missing_assurance": "boolean",
    "missing_charge": "boolean",
    "missing_margin": "boolean",
    "found_price_is_hourly": "boolean",
    "matched_hours_per_day": "number|null",
    "matched_total_agents": "integer|null",
    "confidence": "string",
    "source": "string"
}}

[CRITICAL_INSTRUCTION]
Always use a period (.) as a decimal separator in 'calculation_formula'. Never use a comma. 
If you see a comma in the PDF (e.g. 17,92), convert it to a dot (17.92). 
The 'calculation_formula' should contain the raw arithmetic steps from the table (e.g. "17.92 + 1.72 + 3.78").
"""

    try:
        print("🤖 Asking Gemini to read FULL PDF...")
        
        contents = [
            types.Content(
                role="user",
                parts=[
                    types.Part.from_bytes(data=file_bytes, mime_type='application/pdf'),
                    types.Part.from_text(text=prompt)
                ]
            )
        ]

        result = None
        models_to_try = ["gemini-2.5-flash", "gemini-2.5-pro"]
        
        # --- ATTEMPT WITH AUTO-RETRY LOOP (Handles 503 / 429 Errors) ---
        for model_name in models_to_try:
            for attempt in range(3):
                try:
                    print(f"🤖 Calling {model_name} (Attempt {attempt+1}/3)...")
                    full_response_text = ""
                    config = types.GenerateContentConfig(
                        response_mime_type="application/json",
                    )
                    
                    time.sleep(2)  # Respect RPM limit
                    for chunk in gemini_client.models.generate_content_stream(
                        model=model_name,
                        contents=contents,
                        config=config,
                    ):
                        full_response_text += (chunk.text or "")
                    
                    # 🛡️ VALIDATION INSIDE TRY BLOCK
                    match = re.search(r'\{.*\}', full_response_text, re.DOTALL)
                    if match:
                        result = json.loads(match.group(0))
                        break
                    else:
                        raise ValueError(f"No JSON found in response: {full_response_text[:50]}...")

                except Exception as e:
                    error_msg = str(e)
                    recoverable = [
                        "503", "429", "500", "504", "Overloaded", "UNAVAILABLE",
                        "Expecting ',' delimiter", "Conversion failed",
                        "No JSON found", "Invalid control character", "Extra data"
                    ]
                    
                    if any(err in error_msg for err in recoverable):
                        wait_sec = (attempt + 1) * 3
                        print(f"⏳ Server busy/Recoverable error ({error_msg[:40]}...). Retrying in {wait_sec}s...")
                        time.sleep(wait_sec)
                    else:
                        print(f"⚠️ Non-recoverable Error: {e}")
                        break
            if result:
                break

        # Logic for Processing Result (Only runs if result is set)
        if result:
            price = result.get("price")
            found_in_text_warning = result.get("found_in_text_warning", False)
            confidence = result.get("confidence", "none")
            is_hourly = result.get("found_price_is_hourly", False)
            
            # Granular Missing Column Flags
            missing_assurance = result.get("missing_assurance", False)
            missing_charge = result.get("missing_charge", False)
            missing_margin = result.get("missing_margin", False)
            
            # --- 🎯 ROLE-MATCHED HOURS OVERRIDE ---
            matched_hours = result.get("matched_hours_per_day")
            if matched_hours:
                print(f"🎯 Role-matched hours found in staffing table: {matched_hours}h "
                      f"(overriding previous value: {hours_per_day})")
                hours_per_day = float(matched_hours)

            effective_hours = float(hours_per_day) if hours_per_day else 8.0

            if price and confidence != "none":
                price = float(price)
                matched_total_agents = result.get("matched_total_agents")

                # --- 🏁 GRAND-MASTER SYMBOLIC ENGINE (The Precision Layer) ---
                if result.get("calculation_formula"):
                    formula_price = safe_math_eval(result["calculation_formula"])
                    if formula_price:
                        # 💎 TRUNCATION SHIELD: Always round down for competitive edge
                        price = int(formula_price * 100) / 100.0
                        print(f"🧮 Symbolic Engine Result: {result['calculation_formula']} = {price}")
                # -----------------------------------------------------------

                # --- 🛑 FORFAIT SAFETY GUARD ---
                if "forfait" in u_type and (is_hourly or result.get("is_calculated")):
                    print(f"🛑 Rejecting raw hourly/calculated price ({price}) as a "
                          f"'forfait' total — a forfait must cover the whole "
                          f"contract, not one hour. Falling through instead of "
                          f"submitting a near-zero total.")
                    return None, matched_total_agents

                # --- PYTHON MATH LOGIC (SMART UNIT CONVERSION) ---
                if is_hourly or result.get("is_calculated"):
                    if "jour" in u_type:
                        calculated_price = round(price * effective_hours, 2)
                        print(f"🧮 Python Logic: Converted Hourly ({price}) -> Daily ({calculated_price}) using {effective_hours} hours.")
                        price = calculated_price
                    elif "mois" in u_type:
                        calculated_price = round(price * effective_hours * 26, 2)
                        print(f"🧮 Python Logic: Converted Hourly ({price}) -> Monthly ({calculated_price}) using {effective_hours}h x 26j.")
                        price = calculated_price
                    elif "heure" in u_type:
                        price = round(price, 2)
                else:
                    if "heure" in u_type:
                        calculated_price = round(price / effective_hours, 2)
                        print(f"🧮 Python Logic: Converted Daily ({price}) -> Hourly ({calculated_price}) using {effective_hours} hours.")
                        price = calculated_price
                # -------------------------------------------------
                
                if found_in_text_warning:
                     # Text Warning: Return EXACT value (Priority 1)
                     print(f"👁️ Vision Extracted via WARNING: {price} (Exact Value - No Margin)")
                
                else: 
                    # Apply User's Specific Formula (Context-Aware Adder):
                    # Assurance (+0.01) | charge (+0.01) | Margin (+0.01)
                    adder = 0
                    log_parts = []
                    
                    if missing_assurance: 
                        adder += 0.01
                        log_parts.append("Assur(+0.01)")
                    if missing_charge: 
                        adder += 0.01 
                        log_parts.append("Chg(+0.01)")
                    if missing_margin: 
                        adder += 0.01
                        log_parts.append("Mrg(+0.01)")
                    
                    if adder > 0:
                        price = round(price + adder, 2)
                        print(f"👁️ Vision Adjusted: {price} [{' '.join(log_parts)}]")
                    else:
                        print(f"👁️ Vision Extracted: {price}")
                        
                return price, matched_total_agents
            else:
                 # Check for calculation flag
                if result.get("is_calculated"):
                    print("👁️ Vision calculated price from SMIG formula.")

                matched_total_agents = result.get("matched_total_agents")

                # --- 🎯 ROLE-MATCHED HOURS FALLBACK (using shared SMIG calculator) ---
                if matched_hours or hours_per_day:
                    h = float(matched_hours if matched_hours else hours_per_day)
                    computed_price = compute_smig_unit_price(u_type, effective_hours=h)
                    if computed_price is not None:
                        return computed_price, matched_total_agents

                print(f"👁️ Vision found no price. Confidence: {confidence}")
                return None, matched_total_agents
        else:
            return None, None

    except Exception as e:
        print(f"⚠️ Gemini Vision API Error: {e}")
        return None, None

def read_pdf_from_bytes(file_bytes, unit_type=None, hours_per_day=None, role_description=None):
    """
    Smart PDF Reader (Vision Upgrade).
    1. ZERO-TOKEN PROTOCOL: Quickly scan for keywords or check if Scanned.
    2. VISION GATEKEEPER:
       - If Keywords Found (Score > 0) -> VALID -> Send to Gemini Vision.
       - If No Text (Scanned) -> VALID -> Send to Gemini Vision.
       - If Text Found but No Keywords -> GARBAGE -> Skip.
    3. EXECUTION:
       - Calls extract_price_with_gemini_vision(file_bytes)
       - Returns PRICE (float) or None.
    """
    try:
        # Full unconditional scan as user requested
        print(f"🚀 Vision Engine: ACCELERATED [User Requested 100% Full Scan Mode]. Sending to Gemini...")
        return extract_price_with_gemini_vision(file_bytes, unit_type, hours_per_day, role_description)

    except Exception as e:
        print(f"⚠️ PDF Read Error: {e}")
        return None, None

# --- EXCEL READING (openpyxl) ---
def read_excel_from_bytes(file_bytes):
    """Read Excel content from bytes using openpyxl."""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(filename=BytesIO(file_bytes), data_only=True)
        text = ""
        for sheet in wb.worksheets:
            for row in sheet.iter_rows():
                row_text = " | ".join([str(cell.value) if cell.value else "" for cell in row])
                if row_text.strip():
                    text += row_text + "\n"
        wb.close()
        return text
    except ImportError:
        print("⚠️ openpyxl not installed. Run: pip install openpyxl")
        return None
    except Exception as e:
        print(f"⚠️ Excel read error: {e}")
        return None

# --- WORD READING (python-docx) ---
def read_word_from_bytes(file_bytes):
    """Read Word document content from bytes using python-docx."""
    try:
        from docx import Document
        doc = Document(BytesIO(file_bytes))
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        # Also read tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                text += row_text + "\n"
        return text
    except ImportError:
        print("⚠️ python-docx not installed. Run: pip install python-docx")
        return None
    except Exception as e:
        print(f"⚠️ Word read error: {e}")
        return None

# --- LIBREOFFICE CONVERSION (HEADLESS) ---
def convert_office_to_pdf(file_bytes, input_filename):
    """
    Converts Word/Excel bytes to PDF bytes using LibreOffice (Headless).
    Uses a temporary directory to handle the conversion without cluttering.
    Returns: PDF bytes or None.
    """
    if not os.path.exists(LIBREOFFICE_PATH):
        print("⚠️ LibreOffice not found at configured path. Skipping conversion.")
        return None

    temp_dir = tempfile.mkdtemp()
    
    try:
        # 1. Write Input File
        ext = os.path.splitext(input_filename)[1]
        input_path = os.path.join(temp_dir, f"input{ext}")
        with open(input_path, "wb") as f:
            f.write(file_bytes)
            
        # 2. Run LibreOffice (Headless)
        # Command: soffice --headless --convert-to pdf --outdir <dir> <file>
        cmd = [
            LIBREOFFICE_PATH,
            "--headless",
            "--convert-to", "pdf",
            "--outdir", temp_dir,
            input_path
        ]
        
        # Run with timeout to prevent hanging
        subprocess.run(cmd, check=True, timeout=20, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        
        # 3. Read Output PDF
        pdf_filename = f"input.pdf"
        pdf_path = os.path.join(temp_dir, pdf_filename)
        
        start_time = time.time()
        while time.time() - start_time < 2:
             if os.path.exists(pdf_path):
                 break
             time.sleep(0.1)
             
        if os.path.exists(pdf_path):
            with open(pdf_path, "rb") as f:
                pdf_bytes = f.read()
            print(f"🔄 Converted {input_filename} to PDF ({len(pdf_bytes)} bytes)")
            return pdf_bytes
        else:
            print("⚠️ Conversion failed: PDF output not found.")
            return None
            
    except subprocess.TimeoutExpired:
        print("⚠️ LibreOffice timed out.")
        return None
    except Exception as e:
        print(f"⚠️ Conversion error: {e}")
        return None
    finally:
        # 4. Cleanup (Always)
        try:
            for f in os.listdir(temp_dir):
                try:
                    os.remove(os.path.join(temp_dir, f))
                except: pass
            os.rmdir(temp_dir)
        except:
            pass

# --- FILE READER DISPATCHER ---
def read_file_from_bytes(filename, file_bytes, unit_type=None, hours_per_day=None, role_description=None):
    """Read file content based on extension."""
    filename_lower = filename.lower()
    
    if filename_lower.endswith('.pdf'):
        return read_pdf_from_bytes(file_bytes, unit_type, hours_per_day, role_description)
    
    # 🆕 TRY CONVERSION FOR OFFICE FILES
    elif filename_lower.endswith(('.xlsx', '.xls', '.docx', '.doc')):
        print(f"🔄 Attempting Vision Conversion for: {filename}")
        pdf_bytes = convert_office_to_pdf(file_bytes, filename)
        
        if pdf_bytes:
            print("👁️ Sending Converted PDF to Gemini Vision...")
            return read_pdf_from_bytes(pdf_bytes, unit_type, hours_per_day, role_description)
        else:
            print("⚠️ Conversion failed, falling back to legacy Text Reader.")
            # FALLBACK to Legacy Text Readers
            if filename_lower.endswith(('.xlsx', '.xls')):
                return read_excel_from_bytes(file_bytes)
            elif filename_lower.endswith(('.docx', '.doc')):
                return read_word_from_bytes(file_bytes)

    else:
        # Try to read as text
        try:
            return file_bytes.decode('utf-8', errors='ignore')
        except:
            return None

# --- PRIORITY FILE DETECTION ---
def get_priority_file(file_list):
    """
    Find priority file containing 'Sous détail'.
    Returns (filename, priority_level) or (None, 0) if no priority file.
    """
    for filename in file_list:
        filename_lower = filename.lower()
        if "sous détail des prix" in filename_lower or "sous detail des prix" in filename_lower:
            return filename, 1  # Highest priority
        elif "sous détail" in filename_lower or "sous detail" in filename_lower:
            return filename, 2  # Second priority
    return None, 0

# --- FIND PRIORITY ZIP LINK ---
def find_priority_zip_link(page):
    """
    Find the priority ZIP link on the page.
    Priority: 'Sous détail des prix' > 'Sous détail' > any other ZIP
    Returns (link_element, link_text, href) or (None, None, None)
    """
    try:
        # Find the attachments section
        attachments_section = page.locator("span:has-text('Pièces jointes')").locator("..")
        zip_links = attachments_section.locator("a[href*='download']").all()
        
        if not zip_links:
            print("📎 No ZIP files found in attachments section")
            return None, None, None
        
        # Priority 1: "Sous détail des prix"
        for link in zip_links:
            link_text = link.inner_text().lower()
            if "sous détail des prix" in link_text or "sous detail des prix" in link_text:
                href = link.get_attribute("href")
                print(f"📦 Found priority ZIP: {link.inner_text()}")
                return link, link.inner_text(), href
        
        # Priority 2: "Sous détail"
        for link in zip_links:
            link_text = link.inner_text().lower()
            if "sous détail" in link_text or "sous detail" in link_text:
                href = link.get_attribute("href")
                print(f"📦 Found priority ZIP: {link.inner_text()}")
                return link, link.inner_text(), href
        
        # Priority 3: First available ZIP
        first_link = zip_links[0]
        href = first_link.get_attribute("href")
        print(f"📦 Using first ZIP: {first_link.inner_text()}")
        return first_link, first_link.inner_text(), href
        
    except Exception as e:
        print(f"⚠️ Error finding ZIP links: {e}")
        return None, None, None

# --- DOWNLOAD ZIP FILE ---
def download_zip_file(page, link_element, filename):
    """
    Download ZIP file using Playwright's download handler.
    Returns path to downloaded file or None.
    """
    try:
        # Clean filename
        safe_filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
        if not safe_filename.endswith('.zip'):
            safe_filename += '.zip'
        
        download_path = os.path.join(DOWNLOAD_PATH, safe_filename)
        
        # Use Playwright download
        with page.expect_download() as download_info:
            link_element.click()
        
        download = download_info.value
        download.save_as(download_path)
        
        print(f"✅ Downloaded: {download_path}")
        return download_path
        
    except Exception as e:
        print(f"⚠️ Download error: {e}")
        return None

# --- READ ZIP AND EXTRACT PRICE ---
def read_zip_and_extract_price(zip_path, unit_type, hours_per_day=None, role_description=None):
    """
    Read ZIP file, find priority file, extract text, get price with AI.
    Reads files in memory without extracting to disk.
    """
    try:
        with zipfile.ZipFile(zip_path, 'r') as zip_file:
            file_list = zip_file.namelist()
            
            # Filter out folders and hidden files (handle nested folder structure)
            file_list = [f for f in file_list 
                        if not f.endswith('/') and          # Skip folder entries
                        '__MACOSX' not in f and              # Skip macOS hidden files
                        not f.startswith('._')]              # Skip macOS resource forks
            
            print(f"📂 ZIP contains {len(file_list)} files")
            
            # Find priority file (works with nested paths like "Folder/file.pdf")
            priority_file, priority_level = get_priority_file(file_list)
            
            if priority_file:
                print(f"🎯 Priority file found: {priority_file} (level {priority_level})")
                files_to_read = [priority_file]
            else:
                # Read all supported files (including those in subfolders)
                supported_extensions = ('.pdf', '.xlsx', '.xls', '.docx', '.doc')
                files_to_read = [f for f in file_list if f.lower().endswith(supported_extensions)]
                print(f"📑 Will check {len(files_to_read)} files")
            
            # Read and extract from each file
            for filename in files_to_read:
                print(f"📖 Reading: {filename}")
                
                file_bytes = zip_file.read(filename)
                
                # Check for Vision Price from PDF (returns a (price, matched_total_agents)
                # tuple) or legacy text content (returns a plain str or None).
                result = read_file_from_bytes(filename, file_bytes, unit_type, hours_per_day, role_description)
                
                if isinstance(result, tuple):
                    price, matched_total_agents = result
                    if price is not None:
                        print(f"🎯 Vision Price Found: {price}")
                        return str(price), matched_total_agents
                    if matched_total_agents:
                        # No price, but we DID learn the real total agent count
                        # from the PDF's staffing table — surface it anyway so
                        # the caller can decide whether a flat fallback guess
                        # would be safe (single position) or dangerous (large
                        # multi-agent team), instead of blindly assuming 1.
                        print(f"   ⚠️ No price found in this file, but total agent "
                              f"count was: {matched_total_agents}")
                        return None, matched_total_agents
                    print(f"   ⚠️ File has insufficient text content or Conversion failed.")
                    continue
                
                # Fallback for Text-based files (Legacy Text Reader)
                if isinstance(result, str):
                     print(f"   ⚠️ Raw text returned but Text AI is disabled. Skipping.")
                else:
                    print(f"   ⚠️ File has insufficient text content or Conversion failed.")
            
            print("📄 No price found in any ZIP file")
            return None, None
            
    except zipfile.BadZipFile:
        print(f"⚠️ Invalid ZIP file: {zip_path}")
        return None, None
    except Exception as e:
        print(f"⚠️ ZIP read error: {e}")
        return None, None

# --- MAIN ENTRY POINT ---
def get_price_from_zip(page, unit_type, hours_per_day=None, role_description=None):
    """
    Main entry point for PDF price extraction.
    Called by main script when no price found in description.
    """
    print("\n📦 === PDF EXTRACTION MODE ===")
    
    # Step 1: Find ZIP link
    link_element, link_text, href = find_priority_zip_link(page)
    
    if not link_element:
        print("❌ No ZIP file to download")
        return None, None
    
    # Step 2: Download ZIP
    zip_path = download_zip_file(page, link_element, link_text)
    
    if not zip_path:
        print("❌ ZIP download failed")
        return None, None
    
    # Step 3: Read ZIP and extract price
    price, matched_total_agents = read_zip_and_extract_price(zip_path, unit_type, hours_per_day, role_description)
    
    # Step 4: Cleanup (optional - keep file for debugging)
    if os.path.exists(zip_path):
        try:
            os.remove(zip_path)
        except:
            pass
    
    if price:
        print(f"✅ PDF Price: {price}")
        return str(price), matched_total_agents
    else:
        # Fallback: RETURN NONE so Main Script handles the Hard Fallback
        print("🧮 No price in PDF - Returning None (Triggering Main Script Fallback)")
        return None, matched_total_agents

# --- TEST FUNCTION ---
if __name__ == "__main__":
    print("PDF Price Extractor Module")
    print("This module is designed to be imported by automate marchepublics - Copy.py")
